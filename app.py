import re
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt
from markdown_it import MarkdownIt
from starlette.applications import Starlette
from starlette.concurrency import run_in_threadpool
from starlette.requests import Request
from starlette.responses import JSONResponse, Response
from starlette.staticfiles import StaticFiles
from starlette.templating import Jinja2Templates

from job_store import JobStoreError, job_store
from research_crew import relevel_notes, run_research

LEVELS = ("easy", "medium", "high")

BASE_DIR = Path(__file__).resolve().parent

templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
markdown = MarkdownIt("commonmark", {"html": False, "linkify": True}).enable("table")
executor = ThreadPoolExecutor(max_workers=1)


def _slugify(value):
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return slug[:48] or "research-report"


def _job_payload(job_id):
    job = job_store.get_job(job_id)
    if not job:
        return None

    payload = {
        "id": job["_id"],
        "topic": job["topic"],
        "status": job["status"],
        "stage": job.get("stage"),
        "progress": job.get("progress", 0),
        "stages": _serialize_stages(job.get("stages", [])),
        "message": job["message"],
        "level": job.get("active_level", "medium"),
        "available_levels": list((job.get("versions") or {}).keys()),
        "created_at": _serialize_datetime(job.get("created_at")),
        "updated_at": _serialize_datetime(job.get("updated_at")),
    }

    report_markdown = _active_markdown(job)
    if report_markdown:
        payload["markdown"] = report_markdown
        payload["html"] = markdown.render(report_markdown)
        payload["download_url"] = f"/api/reports/{job_id}/download"

    if job.get("error"):
        payload["error"] = job["error"]

    return payload


def _active_markdown(job):
    """Markdown for the active level: stored in the doc now, on disk for legacy notes."""
    versions = job.get("versions") or {}
    active = versions.get(job.get("active_level", "medium")) or {}
    if active.get("markdown"):
        return active["markdown"]
    if job.get("markdown"):
        return job["markdown"]
    # Legacy notes generated before markdown moved into the document.
    return _read_text_file(active.get("markdown_path") or job.get("markdown_path"))


def _serialize_datetime(value):
    if isinstance(value, datetime):
        return value.isoformat()
    return value


def _serialize_stages(stages):
    serialized = []
    for stage in stages:
        serialized.append(
            {
                **stage,
                "started_at": _serialize_datetime(stage.get("started_at")),
                "ended_at": _serialize_datetime(stage.get("ended_at")),
            }
        )
    return serialized


def _read_text_file(path):
    if not path:
        return None

    report_path = Path(path)
    if not report_path.exists():
        return None

    return report_path.read_text(encoding="utf-8")


def _remove_job_files(job):
    paths = set()
    for key in ("markdown_path", "docx_path"):
        if job.get(key):
            paths.add(job[key])
    for version in (job.get("versions") or {}).values():
        for key in ("markdown_path", "docx_path"):
            if version.get(key):
                paths.add(version[key])

    for path in paths:
        try:
            Path(path).unlink(missing_ok=True)
        except OSError:
            pass


def _add_markdown_line(document, line):
    text = line.strip()
    if not text:
        return

    heading = re.match(r"^(#{1,6})\s+(.+)$", text)
    if heading:
        level = min(len(heading.group(1)), 4)
        document.add_heading(heading.group(2), level=level)
        return

    bullet = re.match(r"^[-*]\s+(.+)$", text)
    if bullet:
        document.add_paragraph(bullet.group(1), style="List Bullet")
        return

    numbered = re.match(r"^\d+\.\s+(.+)$", text)
    if numbered:
        document.add_paragraph(numbered.group(1), style="List Number")
        return

    document.add_paragraph(text)


def _markdown_to_docx(markdown_text, topic, output_path):
    document = Document()
    document.core_properties.title = f"Study Notes: {topic}"
    document.core_properties.subject = topic

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    document.add_heading(f"Study Notes: {topic}", 0)
    document.add_paragraph(datetime.now().strftime("Generated on %B %d, %Y at %I:%M %p"))

    in_code_block = False
    for line in markdown_text.splitlines():
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            document.add_paragraph(line, style="Intense Quote")
            continue

        _add_markdown_line(document, line)

    document.save(output_path)


def _run_job(job_id):
    job = job_store.get_job(job_id)
    if not job:
        return

    try:
        job_store.update_stage(
            job_id,
            stage="research",
            stage_status="running",
            progress=25,
            message="Researcher is searching, comparing sources, and collecting evidence.",
        )
        report_markdown = run_research(job["topic"])

        job_store.update_stage(
            job_id,
            stage="writing",
            stage_status="running",
            progress=75,
            message="Writer is structuring findings into polished study notes.",
        )

        job_store.update_stage(
            job_id,
            stage="complete",
            stage_status="done",
            progress=100,
            message="Notes ready. Read them here or download as Word/PDF.",
            overall_status="complete",
            markdown=report_markdown,
            active_level="medium",
            versions={"medium": {"markdown": report_markdown}},
            error=None,
        )
    except Exception as exc:
        current_job = job_store.get_job(job_id)
        failed_stage = current_job.get("stage", "research") if current_job else "research"
        job_store.mark_failed(
            job_id,
            stage=failed_stage,
            message=f"The {failed_stage} stage stopped before the report could be generated.",
            error=str(exc),
        )


async def set_level(request: Request):
    job_id = request.path_params["job_id"]
    data = await request.json()
    level = (data.get("level") or "").strip().lower()

    if level not in LEVELS:
        return JSONResponse({"error": "Level must be easy, medium, or high."}, status_code=400)

    try:
        job = job_store.get_job(job_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not job or job.get("status") != "complete":
        return JSONResponse({"error": "Notes are not ready yet."}, status_code=404)

    versions = job.get("versions") or {}
    # Backfill for legacy notes created before markdown moved into the document.
    if not versions:
        legacy = job.get("markdown") or _read_text_file(job.get("markdown_path"))
        if legacy:
            versions = {"medium": {"markdown": legacy}}

    if level in versions and versions[level].get("markdown"):
        # Already generated — just switch the active version (no LLM call).
        job_store.update_job(
            job_id,
            active_level=level,
            markdown=versions[level]["markdown"],
        )
    else:
        # Rewrite from the original notes: one cheap LLM pass, no web search.
        source = versions.get("medium") or next(iter(versions.values()), None)
        source_md = (source or {}).get("markdown") or _read_text_file(
            (source or {}).get("markdown_path")
        )
        if not source_md:
            return JSONResponse({"error": "Original notes are missing."}, status_code=409)

        try:
            markdown_text = await run_in_threadpool(
                relevel_notes, job["topic"], source_md, level
            )
        except Exception as exc:
            return JSONResponse({"error": f"Could not rewrite notes: {exc}"}, status_code=500)

        versions[level] = {"markdown": markdown_text}
        job_store.update_job(
            job_id,
            active_level=level,
            versions=versions,
            markdown=markdown_text,
        )

    try:
        payload = _job_payload(job_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)
    return JSONResponse(payload)


async def delete_report(request: Request):
    job_id = request.path_params["job_id"]
    try:
        job = job_store.delete_job(job_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not job:
        return JSONResponse({"error": "Report not found."}, status_code=404)

    _remove_job_files(job)
    job_store.delete_stickies_for_note(job_id)
    return JSONResponse({"ok": True})


def _sticky_payload(sticky):
    return {
        "id": sticky["_id"],
        "note_id": sticky.get("note_id"),
        "text": sticky.get("text", ""),
        "color": sticky.get("color", "yellow"),
        "created_at": _serialize_datetime(sticky.get("created_at")),
        "updated_at": _serialize_datetime(sticky.get("updated_at")),
    }


async def homepage(request: Request):
    return templates.TemplateResponse(request, "index.html")


async def create_report(request: Request):
    data = await request.json()
    topic = (data.get("topic") or "").strip()

    if len(topic) < 3:
        return JSONResponse({"error": "Enter a topic with at least 3 characters."}, status_code=400)

    try:
        job = job_store.create_job(topic=topic, slug=_slugify(topic))
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    executor.submit(_run_job, job["_id"])
    try:
        payload = _job_payload(job["_id"])
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    return JSONResponse(payload, status_code=202)


async def list_reports(request: Request):
    try:
        jobs = job_store.list_jobs()
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    return JSONResponse(
        [
            {
                "id": job["_id"],
                "topic": job.get("topic", ""),
                "status": job.get("status", ""),
                "created_at": _serialize_datetime(job.get("created_at")),
            }
            for job in jobs
        ]
    )


async def get_report(request: Request):
    job_id = request.path_params["job_id"]
    try:
        payload = _job_payload(job_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not payload:
        return JSONResponse({"error": "Report not found."}, status_code=404)

    return JSONResponse(payload)


async def download_report(request: Request):
    job_id = request.path_params["job_id"]
    try:
        job = job_store.get_job(job_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not job or job.get("status") != "complete":
        return JSONResponse({"error": "Report is not ready yet."}, status_code=404)

    report_markdown = _active_markdown(job)
    if not report_markdown:
        return JSONResponse({"error": "Notes are not available."}, status_code=404)

    # Build the .docx on the fly from the stored markdown (nothing kept on disk).
    buffer = BytesIO()
    _markdown_to_docx(report_markdown, job["topic"], buffer)
    buffer.seek(0)

    filename = f"{_slugify(job['topic'])}-study-notes.docx"
    return Response(
        buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def list_stickies(request: Request):
    note_id = request.query_params.get("note_id")
    if not note_id:
        return JSONResponse({"error": "note_id is required."}, status_code=400)

    try:
        stickies = job_store.list_stickies(note_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)
    return JSONResponse([_sticky_payload(s) for s in stickies])


async def create_sticky(request: Request):
    data = await request.json()
    text = (data.get("text") or "").strip()
    color = (data.get("color") or "yellow").strip()
    note_id = (data.get("note_id") or "").strip()
    if not note_id:
        return JSONResponse({"error": "Open a note first."}, status_code=400)
    if not text:
        return JSONResponse({"error": "Note text cannot be empty."}, status_code=400)

    try:
        sticky = job_store.create_sticky(text=text, color=color, note_id=note_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)
    return JSONResponse(_sticky_payload(sticky), status_code=201)


async def update_sticky(request: Request):
    sticky_id = request.path_params["sticky_id"]
    data = await request.json()
    text = data.get("text")
    color = data.get("color")
    if text is not None:
        text = text.strip()
        if not text:
            return JSONResponse({"error": "Note text cannot be empty."}, status_code=400)

    try:
        sticky = job_store.update_sticky(sticky_id, text=text, color=color)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not sticky:
        return JSONResponse({"error": "Note not found."}, status_code=404)
    return JSONResponse(_sticky_payload(sticky))


async def delete_sticky(request: Request):
    sticky_id = request.path_params["sticky_id"]
    try:
        deleted = job_store.delete_sticky(sticky_id)
    except JobStoreError as exc:
        return JSONResponse({"error": str(exc)}, status_code=503)

    if not deleted:
        return JSONResponse({"error": "Note not found."}, status_code=404)
    return JSONResponse({"ok": True})


app = Starlette(debug=True)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
app.add_route("/", homepage)
app.add_route("/api/reports", create_report, methods=["POST"])
app.add_route("/api/reports", list_reports, methods=["GET"])
app.add_route("/api/reports/{job_id}", get_report, methods=["GET"])
app.add_route("/api/reports/{job_id}", delete_report, methods=["DELETE"])
app.add_route("/api/reports/{job_id}/level", set_level, methods=["POST"])
app.add_route("/api/reports/{job_id}/download", download_report)
app.add_route("/api/stickies", list_stickies, methods=["GET"])
app.add_route("/api/stickies", create_sticky, methods=["POST"])
app.add_route("/api/stickies/{sticky_id}", update_sticky, methods=["PUT"])
app.add_route("/api/stickies/{sticky_id}", delete_sticky, methods=["DELETE"])
